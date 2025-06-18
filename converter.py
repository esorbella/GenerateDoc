def converter(file_contents):
    # %%
    import pandas as pd
    from docx import Document
    from docx.shared import Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from docx.shared import RGBColor # Import RGBColor for shading
    from docx.oxml import OxmlElement # Import OxmlElement for shading
    from docx.oxml.ns import qn # Import qn for shading
    from docx.shared import Cm # Import Cm for setting height

    from datetime import datetime
    import re
    import numpy as np # Import numpy for isnan

    ## --- Configuration ---

    # Input and Output Files
    EXCEL_FILENAME = "Sample sched.xlsx" # Define EXCEL_FILENAME here
    OUTPUT_FILENAME = "generated_schedule.docx"

    # Excel Column Indices (0-indexed)
    # Adjust these based on your specific Excel file layout
    STUDENT_INFO_COL_INDEX = 0     # Column containing "Information for student:" and "Major and Department:"
    MAJOR_VALUE_COL_INDEX = 1      # Column containing the Major value (next to "Major and Department:")
    COURSE_CRN_COL_INDEX = None    # Column containing Course CRN (will be found dynamically)
    COURSE_CODE_COL_INDEX = None   # Column containing Course Code (will be found dynamically)
    COURSE_INSTRUCTOR_COL_INDEX = None # Column containing Instructor(s) (will be found dynamically)
    COURSE_CREDITS_COL_INDEX = None  # Column containing Credits (will be found dynamically)
    SCHEDULE_DAYS_COL_INDEX = 8    # Column containing the days abbreviation (e.g., I is index 8)
    SCHEDULE_TIME_COL_INDEX = 9    # Column containing the time range (e.g., J is index 9)
    SCHEDULE_COURSE_COL_INDEX = 1 # Column containing the course entry for the schedule grid (e.g., B is index 1)


    # Row Ranges (0-indexed)
    # Adjust these based on your specific Excel file layout
    STUDENT_INFO_SEARCH_ROWS = range(0, 5)      # Rows to search for "Information for student:"
    MAJOR_SEARCH_ROWS = range(0, 15)           # Rows to search for "Major and Department:"
    COURSE_HEADER_SEARCH_ROWS = range(0, 20)   # Rows to search for the course table headers (CRN, Course, etc.)
    SCHEDULE_DATA_ROWS = range(10, 17)        # Rows containing the actual weekly schedule data grid

    # Placeholder Information (if not extracted from Excel)
    PLACEHOLDER_ADVISOR = "ENTER ADVISOR HERE"
    PLACEHOLDER_COMMENTS = "ENTER COMMENTS HERE"
    PLACEHOLDER_ORIENTATION = "ENTER ORIENTATION HERE"

    # Mapping for Schedule Days Abbreviations
    DAY_MAPPING = {
        "M": ["Mon"],
        "T": ["Tue"],
        "W": ["Wed"],
        "R": ["Thu"],
        "F": ["Fri"],
        "MW": ["Mon", "Wed"],
        "TR": ["Tue", "Thu"],
    }

    # Standard Times for the Output Schedule Grid (Table Rows)
    # These define the fixed rows in your output Word document schedule table
    STANDARD_TIMES = [
        "7:55-8:15 am", "8:25-8:50 am", "9:00-9:20 am", "9:30-9:55 am",
        "10:05-10:25 am", "10:35-11:00 am", "11:10-11:30 am", "11:40-12:05 am",
        "12:15-12:35 pm", "12:45-1:10 pm", "1:20-1:40 pm", "1:50-2:15 pm",
        "2:25-2:45 pm", "2:55-3:20 pm", "3:30-3:50 pm", "4:00-4:25 pm",
        "4:35-6:00 pm", "6:10 – 7:35 pm", "7:45 – 8:10 pm"
    ]

    # Standard Days for the Output Schedule Grid (Table Columns)
    # These define the fixed columns in your output Word document schedule table
    STANDARD_DAYS = ["Mon", "Tue", "Wed", "Thu", "Fri"]

    # Regex pattern to find time ranges in the schedule data column
    TIME_PATTERN = re.compile(r"\d{1,2}:\d{2}.*?-\s*\d{1,2}:\d{2}\s*(?:am|pm)?", re.IGNORECASE)

    ## --- Helper Functions ---
    def parse_time(time_str):
        """Parses a time string into a datetime.time object."""
        if not isinstance(time_str, str):
            return None
        time_str = time_str.strip()
        if not time_str:
            return None

        try:
            # Try formats with am/pm first (most common and specific)
            for fmt in ('%I:%M %p', '%I:%M%p', '%H:%M %p', '%H:%M%p'):
                 try:
                      return datetime.strptime(time_str, fmt).time()
                 except ValueError:
                      pass

            # Then try military time (24-hour)
            try:
                return datetime.strptime(time_str, '%H:%M').time()
            except ValueError:
                pass

            # If neither of the above worked, it might be a 12-hour time *without* AM/PM.
            # This is the ambiguous case. We need to be careful here.
            # Let's try parsing it as a 12-hour format without assuming AM/PM yet.
            try:
                 # Use %I:%M to parse, but this needs careful handling of the result
                 time_obj_no_ampm = datetime.strptime(time_str, '%I:%M').time()

                 # If the time string contains a 12 (e.g., "12:15"), this is likely 12 PM if the overall range is PM
                 # If the time string is 1-11 (e.g., "1:20", "11:40"), it's ambiguous.
                 # This parsing logic might need to happen *outside* this function,
                 # where we have context from the full time range string.
                 # For now, we'll return this and rely on the calling code to handle 12-hour times without am/pm context.
                 return time_obj_no_ampm

            except ValueError:
                 pass

            # If all fail, print warning and return None
            print(f"Warning: Could not parse time string '{time_str}' with expected formats.")
            return None
        except Exception as e:
            print(f"Error parsing time string '{time_str}': {e}")
            return None

    def find_column_index(df, keywords, search_rows):
        """Finds the column index that contains all specified keywords within the search rows."""
        for r_idx in search_rows:
            if r_idx < len(df):
                # Corrected: Use df_raw instead of df
                row_values = [str(val).strip() if pd.notna(val) else "" for val in df_raw.iloc[r_idx].tolist()]
                # Check if all keywords are present in this row's values
                if all(keyword in row_values for keyword in keywords):
                    # Return the 0-indexed column indices of the first occurrence of each keyword
                    col_indices = {}
                    for keyword in keywords:
                        try:
                            col_indices[keyword] = row_values.index(keyword)
                        except ValueError:
                             # This should not happen if all(keyword in row_values for keyword in keywords) is true,
                             # but as a fallback
                             print(f"Warning: Keyword '{keyword}' found in row {r_idx+1} but its index could not be determined.")
                             return None, None, None, None # Indicate failure

                    # Return indices in a fixed order for the course columns
                    return col_indices.get('CRN'), col_indices.get('Course'), col_indices.get('Instructor(s)'), col_indices.get('Credits')

        print(f"Warning: Could not find a row containing all keywords {keywords} in rows {search_rows.start+1}-{search_rows.stop+1}.")
        return None, None, None, None # Return None if no header row is found

    def shade_cell(cell, color_hex="D3D3D3"): # Default to light gray
        """Shades a table cell with the specified color."""
        try:
            print(f"Attempting to shade cell with color: {color_hex}") # Debug print
            # Ensure the color_hex is valid (e.g., "RRGGBB")
            if not re.fullmatch(r"[0-9a-fA-F]{6}", color_hex):
                print(f"Warning: Invalid hex color format '{color_hex}'. Using default light gray.")
                color_hex = "D3D3D3"

            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            # Remove existing shading elements if any to avoid duplicates
            for child in tcPr.xpath('./w:shd'):
                 tcPr.remove(child)

            shading = OxmlElement(qn('w:shd'))
            # Set the fill attribute using the qualified name
            shading.set(qn('w:fill'), color_hex)

            tcPr.append(shading)
            print("Shading element added to cell.") # Debug print

        except Exception as e:
            print(f"Error applying shading to cell: {e}")


    ## --- Data Loading and Extraction ---

    df_raw = None # Initialize df_raw to None
    print(f"Loading data from: {EXCEL_FILENAME}")
    try:
        # Load Excel file without a header initially to search for information rows
        df_raw = pd.read_excel(file_contents, header=None)
        print("Successfully loaded Excel file.")
        print("First 20 rows and 12 columns of the raw DataFrame:")
        # Ensure we don't try to print columns beyond the DataFrame's actual columns
        print(df_raw.head(20).iloc[:, :min(12, df_raw.shape[1])])
        print("-" * 30)

    except FileNotFoundError:
        print(f"Error: The file '{EXCEL_FILENAME}' was not found.")
        # Exit or set a flag to skip subsequent processing if file not found
        # For now, the rest of the script will handle df_raw being None
        pass # Continue execution to allow checking df_raw later
    except Exception as e:
        print(f"An error occurred while loading the Excel file: {e}")
        pass # Continue execution to allow checking df_raw later


    # Check if df_raw was successfully loaded before proceeding
    if df_raw is not None:

        # Extract name and UID
        info_line = None
        # Check within the specified search rows for the info line
        for i in STUDENT_INFO_SEARCH_ROWS:
            if i < len(df_raw) and STUDENT_INFO_COL_INDEX < df_raw.shape[1]:
                cell_value = df_raw.iloc[i, STUDENT_INFO_COL_INDEX]
                if pd.notna(cell_value) and "Information for student:" in str(cell_value):
                    info_line = str(cell_value)
                    break
            elif i >= len(df_raw):
                break # Stop if we exceed the DataFrame rows


        name = ""
        uid = ""
        if info_line:
            match = re.search(r"Information for student:\s*(.*?)\s+\((U\d+)\)", info_line)
            if match:
                name = match.group(1)
                uid = match.group(2)
            else:
                 print(f"Warning: Found info line '{info_line}' but could not parse name and UID.")
        else:
             print(f"Warning: 'Information for student:' line not found in column {STUDENT_INFO_COL_INDEX+1} within rows {STUDENT_INFO_SEARCH_ROWS.start+1}-{STUDENT_INFO_SEARCH_ROWS.stop+1}.")


        # Extract major
        major_line_found = False
        major = ""
        # Check within the specified search rows for the major line
        for i in MAJOR_SEARCH_ROWS:
            if i < len(df_raw) and STUDENT_INFO_COL_INDEX < df_raw.shape[1]:
                cell_value = df_raw.iloc[i, STUDENT_INFO_COL_INDEX]
                if pd.notna(cell_value) and str(cell_value).strip() == "Major and Department:":
                    major_line_found = True
                    # Check if the major value column exists and the cell is not empty
                    if MAJOR_VALUE_COL_INDEX < df_raw.shape[1]:
                         major_value = df_raw.iloc[i, MAJOR_VALUE_COL_INDEX]
                         major = str(major_value).strip() if pd.notna(major_value) else ""
                    else:
                         print(f"Warning: Major value column index {MAJOR_VALUE_COL_INDEX+1} is out of bounds.")
                    break # Found the major line

            elif i >= len(df_raw):
                 break # Stop if we exceed the DataFrame rows

        if not major_line_found:
             print(f"Warning: 'Major and Department:' line not found in column {STUDENT_INFO_COL_INDEX+1} within rows {MAJOR_SEARCH_ROWS.start+1}-{MAJOR_SEARCH_ROWS.stop+1}.")


        # Use placeholder fields
        advisor = PLACEHOLDER_ADVISOR
        comments = PLACEHOLDER_COMMENTS
        orientation = PLACEHOLDER_ORIENTATION

        # Find the course table header row and column indices dynamically
        course_header_keywords = ['CRN', 'Course', 'Instructor(s)', 'Credits']
        COURSE_CRN_COL_INDEX, COURSE_CODE_COL_INDEX, COURSE_INSTRUCTOR_COL_INDEX, COURSE_CREDITS_COL_INDEX = find_column_index(
            df_raw, course_header_keywords, COURSE_HEADER_SEARCH_ROWS
        )

        header_row = []
        course_data = pd.DataFrame()
        course_data_start_index = None

        if COURSE_CRN_COL_INDEX is not None and COURSE_CODE_COL_INDEX is not None:
            # Assuming the header row is the one containing 'CRN' and 'Course'
            # We need to re-find the exact row index where these headers are located
            temp_header_row_index = None
            for i in COURSE_HEADER_SEARCH_ROWS:
                if i < len(df_raw):
                    row_values = [str(val).strip() if pd.notna(val) else "" for val in df_raw.iloc[i].tolist()]
                    if 'CRN' in row_values and 'Course' in row_values:
                         temp_header_row_index = i
                         break
                elif i >= len(df_raw):
                     break


            if temp_header_row_index is not None:
                header_row = [str(col).strip() if pd.notna(col) else "" for col in df_raw.iloc[temp_header_row_index].tolist()]
                course_data_start_index = temp_header_row_index + 1

                # Extract course data starting from the row after the header row
                if course_data_start_index < len(df_raw):
                     course_data = df_raw.iloc[course_data_start_index:].copy()

                     # Ensure the number of columns matches before assigning headers
                     if len(header_row) == course_data.shape[1]:
                         course_data.columns = header_row
                         # pandas automatically handles duplicate column names by appending .1, .2 etc.
                     else:
                         print(f"Warning: Number of columns in header row ({len(header_row)}) does not match data columns ({course_data.shape[1]}). Cannot assign headers.")
                         course_data = pd.DataFrame() # Reset course_data if headers can't be assigned
                else:
                     print(f"Warning: Course data start index {course_data_start_index+1} is beyond the end of the DataFrame. No course data extracted.")
                     course_data = pd.DataFrame() # No data to extract
            else:
                print("Warning: Course header row not found. Cannot extract course data.")
                course_data = pd.DataFrame() # No header means no data


        # *** THIS IS THE SECTION TO FOCUS ON FOR REMOVING EMPTY ROWS FROM THE DATAFRAME ***
        # Drop or filter out rows that don't represent actual courses
        if not course_data.empty:
            # Use the dynamically found column names for filtering
            crn_col_name = 'CRN' if 'CRN' in course_data.columns else None
            course_col_name = 'Course' if 'Course' in course_data.columns else None

            initial_rows = len(course_data)

            if crn_col_name and course_col_name:
                # Filter: Keep rows where BOTH the 'CRN' column is NOT null AND the 'Course' column is NOT null
                # Also ensure the values are not just empty strings after stripping whitespace
                course_data_filtered = course_data[
                    course_data[crn_col_name].apply(lambda x: pd.notna(x) and str(x).strip() != "") &
                    course_data[course_col_name].apply(lambda x: pd.notna(x) and str(x).strip() != "")
                ].copy()

                rows_after_filter = len(course_data_filtered)
                print(f"Filtered course_data: Started with {initial_rows} rows, filtered {initial_rows - rows_after_filter}, ending with {rows_after_filter} rows.")
                course_data = course_data_filtered # Update course_data to the filtered version
            else:
                print("Warning: 'CRN' or 'Course' columns not found after header assignment. Cannot filter rows based on these columns. Proceeding with potentially unfiltered data.")
                # In this case, course_data remains as extracted, potentially containing empty rows


        # Parse courses from the filtered course_data DataFrame
        courses = []
        if not course_data.empty:
            # Use the dynamically found column names if they exist, otherwise use default or placeholder
            crn_col_name = 'CRN' if 'CRN' in course_data.columns else None
            course_col_name = 'Course' if 'Course' in course_data.columns else None
            instructor_col_name = 'Instructor(s)' if 'Instructor(s)' in course_data.columns else None
            credits_col_name = 'Credits' if 'Credits' in course_data.columns else None

            if crn_col_name and course_col_name: # Require CRN and Course columns to exist for parsing
                for _, row in course_data.iterrows():
                    # Get values using .get() with default "" for robustness
                    crn = str(row.get(crn_col_name, "")).strip()
                    course_code = str(row.get(course_col_name, "")).strip()
                    instructor = str(row.get(instructor_col_name, "")).strip() if instructor_col_name else ""
                    credits_raw = row.get(credits_col_name) if credits_col_name else 0

                    credits = 0
                    if pd.notna(credits_raw):
                         try:
                             # Safely convert credits, handling potential issues
                             credits_str = str(credits_raw).replace(',', '').strip()
                             if credits_str:
                                 credits = int(float(credits_str)) # Use float first for robustness
                         except (ValueError, TypeError):
                             credits = 0 # Set to 0 if conversion fails

                    # Add the course only if CRN and Course Code are present after stripping
                    # This check is technically redundant if the filtering above works correctly,
                    # but it adds an extra layer of safety for creating the 'courses' list.
                    if crn and course_code:
                        courses.append({
                            "CRN": crn,
                            "Course": course_code,
                            "Instructor": instructor,
                            "Credits": str(credits) # Store as string for consistency
                        })
            else:
                print("Warning: 'CRN' or 'Course' columns not found in filtered course data. Cannot parse course list.")
        else:
            print("No course data found after filtering.")

        # --- Debugging: Print the courses list before generating the Word table ---
        print("\n--- Contents of 'courses' list before Word table generation ---")
        print(f"Number of courses in list: {len(courses)}")
        for i, course in enumerate(courses):
            print(f"  Course {i+1}: {course}")
        print("-----------------------------------------------------------\n")


        # Calculate total credits
        total_credits = 0
        if courses:
            try:
                # Sum credits, ensuring they are digits
                total_credits = sum(int(c.get("Credits", "0")) for c in courses if c.get("Credits", "0").isdigit())
            except ValueError:
                print("Warning: Could not calculate total credits due to non-integer credit values.")
                total_credits = "Error calculating"


        ## --- Schedule Grid Parsing ---

        # Create an empty target schedule grid
        schedule_grid_data = {time: {day: "" for day in STANDARD_DAYS} for time in STANDARD_TIMES}

        # Ensure required columns exist for schedule parsing
        if SCHEDULE_DAYS_COL_INDEX < df_raw.shape[1] and SCHEDULE_TIME_COL_INDEX < df_raw.shape[1] and SCHEDULE_COURSE_COL_INDEX < df_raw.shape[1]:
            # Ensure the SCHEDULE_DATA_ROWS range is valid for the DataFrame
            valid_schedule_rows = range(
                SCHEDULE_DATA_ROWS.start,
                min(SCHEDULE_DATA_ROWS.stop, len(df_raw))
            )
            if not valid_schedule_rows:
                 print("Warning: No valid rows found within the specified SCHEDULE_DATA_ROWS range.")
            else:
                print(f"\nParsing schedule data from rows {valid_schedule_rows.start+1}-{valid_schedule_rows.stop} using columns I ({SCHEDULE_DAYS_COL_INDEX+1}), J ({SCHEDULE_TIME_COL_INDEX+1}), and B ({SCHEDULE_COURSE_COL_INDEX+1})...")

                # Iterate through the specified rows that contain schedule entries
                for r_idx in valid_schedule_rows:
                    # Get raw values from the defined schedule columns
                    raw_time_range_value = str(df_raw.iloc[r_idx, SCHEDULE_TIME_COL_INDEX]).strip()
                    raw_days_value = str(df_raw.iloc[r_idx, SCHEDULE_DAYS_COL_INDEX]).strip()
                    raw_course_entry = str(df_raw.iloc[r_idx, SCHEDULE_COURSE_COL_INDEX]).strip() # Changed variable name for clarity

                    # --- Always print the raw data for the row being processed for debugging ---
                    print(f"  Processing row {r_idx+1} (Index {r_idx}): Time='{raw_time_range_value}', Days='{raw_days_value}', Course='{raw_course_entry}'")

                    # Only process if we have time, days, and a course entry
                    if raw_time_range_value and raw_days_value and raw_course_entry:
                         # --- Parse the raw time range ---
                         start_time_str = None
                         end_time_str = None
                         time_range_match = TIME_PATTERN.search(raw_time_range_value)
                         if time_range_match:
                              parts = time_range_match.group(0).split('-')
                              if len(parts) == 2:
                                  start_time_str = parts[0].strip()
                                  end_time_str = parts[1].strip()
                              else:
                                   print(f"Warning: Could not split time range '{time_range_match.group(0)}' into start and end times.")

                         # Convert extracted time strings to datetime.time objects
                         start_time_obj = parse_time(start_time_str)
                         end_time_obj = parse_time(end_time_str)

                         # --- ADD DEBUG PRINT HERE ---
                         print(f"    Parsed Time: Start='{start_time_str}' ({start_time_obj}), End='{end_time_str}' ({end_time_obj})")
                         # --- END DEBUG PRINT ---


                         # --- Determine which standard time slots are covered by this range ---
                         covered_standard_times = []
                         # Only proceed if both start and end times were successfully parsed
                         if start_time_obj is not None and end_time_obj is not None:
                             for std_time_range in STANDARD_TIMES:
                                  std_parts = std_time_range.split('-')
                                  if len(std_parts) == 2:
                                       std_start_str_raw = std_parts[0].strip()
                                       std_end_str_raw = std_parts[1].strip()

                                       # Determine if the standard range is AM or PM based on the full string or end part
                                       is_pm_range = 'pm' in std_time_range.lower() or ('pm' in std_end_str_raw.lower() and 'am' not in std_start_str_raw.lower())
                                       is_am_range = 'am' in std_time_range.lower() or ('am' in std_end_str_raw.lower() and 'pm' not in std_start_str_raw.lower())

                                       # Attempt to parse start and end times, providing AM/PM context explicitly
                                       std_start_obj = None
                                       std_end_obj = None

                                       # Force AM/PM if not present in the raw string part, based on the range
                                       if 'am' in std_start_str_raw.lower() or 'pm' in std_start_str_raw.lower():
                                           # If AM/PM is already there, parse directly
                                           std_start_obj = parse_time(std_start_str_raw)
                                       elif is_am_range:
                                            # If it's an overall AM range and no AM/PM in start, add " am"
                                            std_start_obj = parse_time(std_start_str_raw + " am")
                                       elif is_pm_range:
                                           # If it's an overall PM range and no AM/PM in start, add " pm"
                                           # This covers 12:xx pm and 1:xx pm etc.
                                           std_start_obj = parse_time(std_start_str_raw + " pm")
                                       else:
                                            # If no AM/PM info in string or range, try parsing as is (might default to AM)
                                            # This is a fallback and still potentially ambiguous
                                            std_start_obj = parse_time(std_start_str_raw)


                                       # Parse end time - apply the same logic
                                       if 'am' in std_end_str_raw.lower() or 'pm' in std_end_str_raw.lower():
                                           std_end_obj = parse_time(std_end_str_raw)
                                       elif is_am_range:
                                            std_end_obj = parse_time(std_end_str_raw + " am")
                                       elif is_pm_range:
                                            std_end_obj = parse_time(std_end_str_raw + " pm")
                                       else:
                                           std_end_obj = parse_time(std_end_str_raw)


                                       # --- Keep your debug prints here ---
                                       print(f"      Checking Standard Time Slot '{std_time_range}': Parsed Start='{std_start_str_raw}' ({std_start_obj}), Parsed End='{std_end_str_raw}' ({std_end_obj})")
                                       print(f"        Comparison: {std_start_obj} >= {start_time_obj} ({std_start_obj >= start_time_obj}) AND {std_end_obj} <= {end_time_obj} ({std_end_obj <= end_time_obj})")
                                       # --- END DEBUG PRINT ---


                                       if std_start_obj is not None and std_end_obj is not None:
                                           # Check if the standard time slot is fully within or exactly matches the raw time range
                                           if std_start_obj >= start_time_obj and std_end_obj <= end_time_obj:
                                               covered_standard_times.append(std_time_range)
                                  else:
                                       # Handle non-standard time strings in STANDARD_TIMES (like "Evening")
                                       if std_time_range.strip().lower() == raw_time_range_value.strip().lower():
                                            covered_standard_times.append(std_time_range)


                         # Parse the raw days string to get standard days
                         days_to_populate = []
                         # Check if the raw days value is a key in our mapping
                         if raw_days_value in DAY_MAPPING:
                              days_to_populate = DAY_MAPPING[raw_days_value]
                         else:
                             # Fallback: Handle single-letter days if not explicitly in mapping
                             if len(raw_days_value) == 1 and raw_days_value.upper() in [d[0] for d in STANDARD_DAYS]:
                                 day_letter = raw_days_value.upper()
                                 for std_day in STANDARD_DAYS:
                                      if std_day.startswith(day_letter):
                                           days_to_populate.append(std_day)
                                           break # Assume unique first letters


                         # --- Populate the schedule grid data if we have covered times and days ---
                         if covered_standard_times and days_to_populate:
                             print(f"  Mapping '{raw_course_entry}' for time range '{raw_time_range_value}' (covers: {covered_standard_times}) on days {days_to_populate}") # Debug print
                             for std_time_slot in covered_standard_times:
                                  for day in days_to_populate:
                                      if std_time_slot in schedule_grid_data and day in schedule_grid_data[std_time_slot]:
                                          # Append the course entry to the correct cell in the grid
                                          # Add a newline if the cell is not empty to separate multiple entries
                                          if schedule_grid_data[std_time_slot][day]:
                                               schedule_grid_data[std_time_slot][day] += "\n" + raw_course_entry
                                          else:
                                               schedule_grid_data[std_time_slot][day] = raw_course_entry
                                      else:
                                           print(f"Warning: Could not map course entry '{raw_course_entry}' to unknown standard time '{std_time_slot}' or day '{day}'. Check STANDARD_TIMES and STANDARD_DAYS configuration.")
                         else:
                              # Print why a row wasn't mapped to help debugging
                              if raw_time_range_value and raw_days_value and raw_course_entry: # Only print if data was present but not mapped
                                 print(f"  Row {r_idx+1} skipped due to missing data in key columns: Time='{raw_time_range_value}', Days='{raw_days_value}', Course='{raw_course_entry}'")

                    else:
                         # This row is skipped because time, days, or course entry was empty/missing
                         if any([pd.notna(df_raw.iloc[r_idx, SCHEDULE_TIME_COL_INDEX]),
                                 pd.notna(df_raw.iloc[r_idx, SCHEDULE_DAYS_COL_INDEX]),
                                 pd.notna(df_raw.iloc[r_idx, SCHEDULE_COURSE_COL_INDEX])]):
                              print(f"  Row {r_idx+1} skipped due to missing data in key columns: Time='{raw_time_range_value}', Days='{raw_days_value}', Course='{raw_course_entry}'")


        else:
             if SCHEDULE_DAYS_COL_INDEX is not None and SCHEDULE_DAYS_COL_INDEX >= df_raw.shape[1]:
                  print(f"Cannot parse schedule data: Days column index {SCHEDULE_DAYS_COL_INDEX+1} is out of bounds for DataFrame with {df_raw.shape[1]} columns.")
             if SCHEDULE_TIME_COL_INDEX is not None and SCHEDULE_TIME_COL_INDEX >= df_raw.shape[1]:
                  print(f"Cannot parse schedule data: Time column index {SCHEDULE_TIME_COL_INDEX+1} is out of bounds for DataFrame with {df_raw.shape[1]} columns.")
             if SCHEDULE_COURSE_COL_INDEX is not None and SCHEDULE_COURSE_COL_INDEX >= df_raw.shape[1]:
                  print(f"Cannot parse schedule data: Schedule Course column index {SCHEDULE_COURSE_COL_INDEX+1} is out of bounds for DataFrame with {df_raw.shape[1]} columns.")

    else: # If df_raw was not successfully loaded
        print("Skipping data extraction and Word document generation due to file loading error.")
        # Initialize necessary variables to default values to prevent errors later if the script didn't exit
        name = ""
        uid = ""
        major = ""
        advisor = PLACEHOLDER_ADVISOR
        comments = PLACEHOLDER_COMMENTS
        orientation = PLACEHOLDER_ORIENTATION
        courses = []
        total_credits = 0
        schedule_grid_data = {time: {day: "" for day in STANDARD_DAYS} for time in STANDARD_TIMES}


    ## --- Generate Word doc ---
    # Only generate the document if the file was loaded, or handle the case where data is empty
    if df_raw is not None:
        doc = Document()

        # --- Section 1: Student Information (Table Format) ---
        doc.add_paragraph("\n\nNew York Campus, Undergraduate\n")

        # Create a 4x2 table for student information
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid' # Optional: Apply a table style

        # Populate the table cells according to the corrected layout
        # Row 0: UID (Col 1), Orientation Date (Col 2)
        uid_cell_col1 = info_table.cell(0, 0)
        uid_paragraph_col1_r0 = uid_cell_col1.paragraphs[0]
        uid_paragraph_col1_r0.text = "" # Clear existing text
        run = uid_paragraph_col1_r0.add_run("UID:")
        run.bold = True
        uid_paragraph_col1_r0.add_run(f" {uid}")

        orientation_cell_col2 = info_table.cell(0, 1)
        orientation_paragraph_col2_r0 = orientation_cell_col2.paragraphs[0]
        orientation_paragraph_col2_r0.text = "" # Clear existing text
        run = orientation_paragraph_col2_r0.add_run("Orientation Date:")
        run.bold = True
        orientation_paragraph_col2_r0.add_run(f" {orientation}")


        # Row 1: Name (Col 1), Total Credits (Col 2)
        name_cell_col1 = info_table.cell(1, 0)
        name_paragraph_col1_r1 = name_cell_col1.paragraphs[0]
        name_paragraph_col1_r1.text = "" # Clear existing text
        run = name_paragraph_col1_r1.add_run("Name:")
        run.bold = True
        name_paragraph_col1_r1.add_run(f" {name}")

        credits_cell_col2 = info_table.cell(1, 1)
        credits_paragraph_col2_r1 = credits_cell_col2.paragraphs[0]
        credits_paragraph_col2_r1.text = "" # Clear existing text
        run = credits_paragraph_col2_r1.add_run("Total Credits:")
        run.bold = True
        credits_paragraph_col2_r1.add_run(f" {total_credits}")


        # Row 2: Major (Col 1), Advisor (Col 2)
        major_cell_col1 = info_table.cell(2, 0)
        major_paragraph_col1_r2 = major_cell_col1.paragraphs[0]
        major_paragraph_col1_r2.text = "" # Clear existing text
        run = major_paragraph_col1_r2.add_run("Major:")
        run.bold = True
        major_paragraph_col1_r2.add_run(f" {major}")

        advisor_cell_col2 = info_table.cell(2, 1)
        advisor_paragraph_col2_r2 = advisor_cell_col2.paragraphs[0]
        advisor_paragraph_col2_r2.text = "" # Clear existing text
        run = advisor_paragraph_col2_r2.add_run("Advisor:")
        run.bold = True
        advisor_paragraph_col2_r2.add_run(f" {advisor}")

        # Row 3: Comments (Col 1), Blank Cell (Col 2)
        comments_cell_col1 = info_table.cell(3, 0)
        comments_paragraph_col1_r3 = comments_cell_col1.paragraphs[0]
        comments_paragraph_col1_r3.text = "" # Clear existing text
        run = comments_paragraph_col1_r3.add_run("Comments:")
        run.bold = True
        comments_paragraph_col1_r3.add_run(f" {comments}")

        info_table.cell(3, 1).text = "" # Keep this cell blank as requested


        # Adjust text alignment and vertical alignment for the info table
        for row in info_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT # Align text to the left
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER # Vertically center text

        # Set minimum row height for the info table
        info_table_row_height_cm = 0.8 # Adjust this value as needed
        for row in info_table.rows:
            row.height = Cm(info_table_row_height_cm)
            row.height_rule = 2 # 2 means at least, 1 means exact, 0 means auto


        # --- Section 2: Course Table ---
        doc.add_paragraph("\n") # Spacing after info table
        doc.add_paragraph("Courses:").bold = True

        if courses: # Check if the 'courses' list is not empty after filtering
            # Add 1 for the header row
            # Ensure the table has exactly len(courses) + 1 rows (header + one row per course)
            table = doc.add_table(rows=len(courses) + 1, cols=5)
            table.style = 'Table Grid'

            hdr = table.rows[0].cells
            hdr[0].text = 'CRN'
            hdr[1].text = 'Course'
            hdr[2].text = 'Instructor'
            hdr[3].text = 'Crdt'
            hdr[4].text = 'Attribute'

            for cell in hdr:
                for paragraph in cell.paragraphs:
                     for run in paragraph.runs:
                         run.bold = True
                     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Iterate over the filtered 'courses' list and populate table rows
            for i, course in enumerate(courses):
                # Use i + 1 to get the correct row index after the header row
                row = table.rows[i + 1].cells
                row[0].text = course.get("CRN")
                row[1].text = course.get("Course")
                row[2].text = course.get("Instructor")
                row[3].text = course.get("Credits")
                row[4].text = "" # Attribute was not extracted, leave blank or add placeholder

                for cell in row:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Set course table column widths
            try:
                # Calculate available page width (roughly)
                total_page_width_inches = (doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin) / 914400.0 # Convert EMUs to inches
                # Assign widths (adjust these proportions as needed)
                col_widths_inches = [1.0, 1.8, 1.8, 0.7, 1.2] # Example widths in inches
                total_assigned_width = sum(col_widths_inches)

                # Scale widths if the sum exceeds the page width (optional, but helps prevent overflow)
                if total_assigned_width > total_page_width_inches:
                    scale_factor = total_page_width_inches / total_assigned_width
                    col_widths_inches = [w * scale_factor for w in col_widths_inches]

                for j in range(min(len(col_widths_inches), len(table.columns))):
                     table.columns[j].width = Inches(col_widths_inches[j])

            except Exception as e:
                 print(f"An error occurred while setting course table column widths: {e}")

            # Set minimum row height for the course table (including header row)
            course_table_row_height_cm = 0.8 # Increased height
            for row in table.rows:
                 row.height = Cm(course_table_row_height_cm)
                 row.height_rule = 2 # 2 means at least, 1 means exact, 0 means auto


        else:
             doc.add_paragraph("No courses found.")


        doc.add_paragraph("\n") # Spacing

        # --- Section 3: Weekly Schedule Grid ---
        doc.add_paragraph("SID:").bold = True # Add "SID:" heading (UID is often called SID)

        # Add 1 to the number of columns for the Time column
        schedule_table = doc.add_table(rows=len(STANDARD_TIMES) + 1, cols=len(STANDARD_DAYS) + 1)
        schedule_table.style = 'Table Grid'
        schedule_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set header row (Time | Mon | Tue | ...)
        header_cells = schedule_table.rows[0].cells
        header_cells[0].text = "Time"
        for d_idx, day in enumerate(STANDARD_DAYS):
            header_cells[d_idx + 1].text = day

        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                     run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


        # Populate the schedule table with data from schedule_grid_data
        for t_idx, time_slot in enumerate(STANDARD_TIMES):
            row_cells = schedule_table.rows[t_idx + 1].cells # +1 to skip header row
            row_cells[0].text = time_slot # First cell is the time slot

            for d_idx, day in enumerate(STANDARD_DAYS):
                # Get the course data for this time and day from our grid
                course_entry = schedule_grid_data.get(time_slot, {}).get(day, "")

                # Check for Common Hour time slots and days
                is_common_hour_time = time_slot in ["12:15-12:35 pm", "12:45-1:10 pm"]
                is_common_hour_day = day in ["Mon", "Wed", "Fri"]

                if is_common_hour_time and is_common_hour_day:
                    row_cells[d_idx + 1].text = "COMMON HOUR" # Set text to COMMON HOUR
                    # shade_cell(row_cells[d_idx + 1], "D3D3D3") # Shade the cell gray - Commented out as requested
                else:
                    row_cells[d_idx + 1].text = course_entry # Populate day cells with extracted course data

                # Center text and vertically align in data cells
                for paragraph in row_cells[d_idx + 1].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[d_idx + 1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Set schedule table column widths
        try:
            total_page_width_inches = (doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin) / 914400.0
            num_cols = len(STANDARD_DAYS) + 1
            if num_cols > 0:
                 # Assign a fixed width to the time column and distribute the rest
                 time_col_width_inches = 1.5
                 if num_cols > 1:
                    schedule_table.columns[0].width = Inches(time_col_width_inches)
                    remaining_width_inches = total_page_width_inches - time_col_width_inches
                    day_column_width_inches = remaining_width_inches / (num_cols - 1)

                    # Ensure day column width is not negative or zero
                    day_column_width_inches = max(0.1, day_column_width_inches)

                    for j in range(1, num_cols):
                        schedule_table.columns[j].width = Inches(day_column_width_inches)
                 elif num_cols == 1:
                     schedule_table.columns[0].width = Inches(total_page_width_inches) # Should not happen with standard days

        except Exception as e:
             print(f"An error occurred while setting schedule table column widths: {e}")

        # Set minimum row height for the schedule grid table
        schedule_table_row_height_cm = 1.0 # Increased height
        for row in schedule_table.rows:
            row.height = Cm(schedule_table_row_height_cm)
            row.height_rule = 2 # 2 means at least, 1 means exact, 0 means auto


        doc.add_paragraph("\n\n") # Spacing


        # --- Section 4: Office Use Only ---
        doc.add_paragraph("OFFICE USE ONLY").bold = True

        # Using paragraphs for simplicity, adjust if sample doc uses a table
        doc.add_paragraph("Creator: _________________ Tally Sheet Updated:____ Hold:_____")
        doc.add_paragraph("M PMT: _________________")
        doc.add_paragraph("Reviser: _________________ No Changes Made: ______ Schedule Changes Made (if Applicable): ______ Tally Sheet Updated: ______")


        ## --- Save the document ---
        doc.save(OUTPUT_FILENAME)
        print(f"\nGenerated Word document: {OUTPUT_FILENAME}")
    else:
        print("\nWord document not generated due to errors in data loading or processing.")